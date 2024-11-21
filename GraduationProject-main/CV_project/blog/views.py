import csv
import os
import PyPDF2
from django.http import HttpResponse
from openpyxl import Workbook
from django.shortcuts import render
from django.core.files.storage import FileSystemStorage
from docx import Document  # Import for handling Word documents

def csv_to_excel(request):
    if request.method == 'POST' and 'csv_file' in request.FILES:
        # Save uploaded file
        csv_file = request.FILES['csv_file']
        fs = FileSystemStorage()
        file_path = fs.save(csv_file.name, csv_file)
        file_path = fs.path(file_path)

        # Check file extension
        allowed_extensions = ['csv', 'tsv', 'xlsx', 'txt', 'pdf', 'docx']
        file_extension = csv_file.name.split('.')[-1].lower()
        if file_extension not in allowed_extensions:
            return HttpResponse("Unsupported file type.", status=400)

        data = []

        # Read data based on file extension
        try:
            if file_extension in ['csv', 'tsv', 'txt']:
                with open(file_path, newline='', encoding='utf-8') as file:
                    reader = csv.reader(file)
                    for row in reader:
                        data.append(row)
            elif file_extension == 'pdf':
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            data.append(text.split('\n'))  # Split text into lines
            elif file_extension == 'docx':
                doc = Document(file_path)
                for para in doc.paragraphs:
                    data.append(para.text.split('\n'))  # Split text into lines
        except Exception as e:
            return HttpResponse(f"Error reading file: {str(e)}", status=400)

        # Create a new Excel workbook
        workbook = Workbook()
        sheet = workbook.active

        # Write data to the Excel sheet
        for row in data:
            sheet.append(row)

        # Save the Excel file to the response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename=output.xlsx'
        workbook.save(response)

        return response

    # Render the file upload page if not a POST request
    return render(request, 'blog/upload_csv.html')

def extract_skills_experience_education_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()

    # Use basic string methods or regex to find relevant sections
    skills = extract_section(content, "Skills")
    experience = extract_section(content, "Professional Experience")
    education = extract_section(content, "Education")

    return skills, experience, education

def extract_section(content, section_title):
    # Find the section in the content
    start = content.lower().find(section_title.lower())
    if start == -1:
        return None
    end = content.find("\n", start + len(section_title))
    return content[start:end].strip()  # Adjust as necessary

def extract_skills_experience_education_pdf(file_path):
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        content = ''
        for page in reader.pages:
            text = page.extract_text()
            if text:
                content += text + "\n"  # Aggregate text from all pages

    return extract_skills_experience_education_txt(content)  # Reuse existing function

def extract_skills_experience_education_docx(file_path):
    doc = Document(file_path)
    content = '\n'.join([para.text for para in doc.paragraphs])
    return extract_skills_experience_education_txt(content)  # Reuse existing function

def extract_cv_data(request):
    if request.method == 'POST' and 'cv_file' in request.FILES:
        cv_file = request.FILES['cv_file']
        fs = FileSystemStorage()
        file_path = fs.save(cv_file.name, cv_file)
        file_path = fs.path(file_path)

        # Determine file type and call the appropriate function
        _, file_extension = os.path.splitext(cv_file.name)
        file_extension = file_extension.lower()

        if file_extension == '.txt':
            skills, experience, education = extract_skills_experience_education_txt(file_path)
        elif file_extension == '.pdf':
            skills, experience, education = extract_skills_experience_education_pdf(file_path)
        elif file_extension == '.docx':
            skills, experience, education = extract_skills_experience_education_docx(file_path)
        else:
            return HttpResponse("Unsupported file type.", status=400)

        # Render or return the extracted data
        return render(request, 'your_template.html', {
            'skills': skills,
            'experience': experience,
            'education': education,
        })

    return render(request, 'upload_cv.html')
